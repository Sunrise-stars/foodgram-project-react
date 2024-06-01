import random
from docx import Document

# Чтение исходного документа
doc = Document(r'C:\Users\User\Desktop\dev\1.docx')
new_doc = Document()

# Генерация случайных цен и добавление их к строкам
for para in doc.paragraphs:
    text = para.text
    # Генерация случайной цены в диапазоне от 80 до 350
    price = random.randint(80, 350)
    new_text = f"{text},{price}"
    new_doc.add_paragraph(new_text)

# Сохранение обновленного документа
output_path = r'C:\Users\User\Desktop\dev\updated_ingredients_with_random_prices.docx'
new_doc.save(output_path)
